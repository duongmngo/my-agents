"""
DOCX file extractor using python-docx
"""
import io
import logging
from pathlib import Path

from app.services.extractors import BaseExtractor, ExtractionResult

logger = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """Extractor for Microsoft Word DOCX files"""
    
    supported_extensions = ['docx']
    
    def extract(self, file_path: str) -> ExtractionResult:
        """Extract text from a DOCX file"""
        try:
            from docx import Document # type: ignore
        except ImportError:
            return ExtractionResult(
                text="",
                error="python-docx library not installed. Run: pip install python-docx"
            )
        
        try:
            path = Path(file_path)
            if not path.exists():
                return ExtractionResult(
                    text="",
                    error=f"File not found: {file_path}"
                )
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            cleaned_text = self._clean_text(full_text)
            
            # Extract metadata from core properties
            metadata = {}
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
            except Exception:
                pass
            
            # Estimate page count (rough approximation)
            # Average ~3000 characters per page
            estimated_pages = max(1, len(cleaned_text) // 3000) if cleaned_text else 0
            
            return ExtractionResult(
                text=cleaned_text,
                page_count=estimated_pages,
                metadata=metadata if metadata else None
            )
            
        except Exception as e:
            logger.exception(f"Error extracting text from DOCX: {file_path}")
            return ExtractionResult(
                text="",
                error=f"Failed to extract text from DOCX: {str(e)}"
            )
    
    def extract_from_bytes(self, content: bytes, filename: str) -> ExtractionResult:
        """Extract text from DOCX bytes content"""
        try:
            from docx import Document # type: ignore
        except ImportError:
            return ExtractionResult(
                text="",
                error="python-docx library not installed. Run: pip install python-docx"
            )
        
        try:
            docx_stream = io.BytesIO(content)
            doc = Document(docx_stream)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            cleaned_text = self._clean_text(full_text)
            
            # Extract metadata
            metadata = {}
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
            except Exception:
                pass
            
            # Estimate page count
            estimated_pages = max(1, len(cleaned_text) // 3000) if cleaned_text else 0
            
            return ExtractionResult(
                text=cleaned_text,
                page_count=estimated_pages,
                metadata=metadata if metadata else None
            )
            
        except Exception as e:
            logger.exception(f"Error extracting text from DOCX bytes: {filename}")
            return ExtractionResult(
                text="",
                error=f"Failed to extract text from DOCX: {str(e)}"
            )
